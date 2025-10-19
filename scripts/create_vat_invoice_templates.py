#!/usr/bin/env python3.11
"""
Create VAT-Aware Invoice Templates for South African SMMEs
Generates SARS-compliant invoice templates in DOCX format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border_to_cell(cell, **kwargs):
    """Add borders to table cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:color'), '1C2A4A')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def create_full_tax_invoice():
    """Create Full Tax Invoice template (for transactions > R5,000)"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header with business details
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    # Left column - Business name
    left_cell = header_table.rows[0].cells[0]
    p = left_cell.paragraphs[0]
    run = p.add_run("YOUR BUSINESS NAME")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 42, 74)  # Navy
    
    p = left_cell.add_paragraph("Your Address")
    p = left_cell.add_paragraph("City, Province, Postal Code")
    p = left_cell.add_paragraph("Phone: +27 XX XXX XXXX")
    p = left_cell.add_paragraph("Email: your@email.com")
    p = left_cell.add_paragraph("VAT No: 4XXXXXXXXX")
    
    # Right column - Invoice title and details
    right_cell = header_table.rows[0].cells[1]
    right_cell.width = Inches(2.5)
    p = right_cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("TAX INVOICE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 42, 74)
    
    # Invoice details table
    details_table = doc.add_table(rows=3, cols=2)
    details_table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    details_table.rows[0].cells[0].text = "Invoice #:"
    details_table.rows[0].cells[1].text = "INV-001"
    details_table.rows[1].cells[0].text = "Date:"
    details_table.rows[1].cells[1].text = "2024/10/18"
    details_table.rows[2].cells[0].text = "Due Date:"
    details_table.rows[2].cells[1].text = "2024/11/17"
    
    for row in details_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    right_cell.add_paragraph()._element.append(details_table._element)
    
    doc.add_paragraph()  # Spacing
    
    # Bill To section
    p = doc.add_paragraph()
    run = p.add_run("BILL TO:")
    run.font.bold = True
    run.font.size = Pt(12)
    
    doc.add_paragraph("Client Name")
    doc.add_paragraph("Client Address")
    doc.add_paragraph("City, Province, Postal Code")
    doc.add_paragraph("VAT No: (if applicable)")
    
    doc.add_paragraph()  # Spacing
    
    # Items table
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = items_table.rows[0].cells
    headers = ['DESCRIPTION', 'QTY', 'UNIT PRICE', 'VAT', 'AMOUNT']
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Set background color to navy
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1C2A4A')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Sample rows
    for i in range(3):
        row_cells = items_table.add_row().cells
        row_cells[0].text = f"Item {i+1} description"
        row_cells[1].text = "1"
        row_cells[2].text = "R 1,000.00"
        row_cells[3].text = "R 150.00"
        row_cells[4].text = "R 1,150.00"
    
    doc.add_paragraph()  # Spacing
    
    # Totals section
    totals_table = doc.add_table(rows=4, cols=2)
    totals_table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    totals_table.rows[0].cells[0].text = "SUBTOTAL:"
    totals_table.rows[0].cells[1].text = "R 3,000.00"
    totals_table.rows[1].cells[0].text = "VAT (15%):"
    totals_table.rows[1].cells[1].text = "R 450.00"
    totals_table.rows[2].cells[0].text = "TOTAL:"
    totals_table.rows[2].cells[1].text = "R 3,450.00"
    totals_table.rows[3].cells[0].text = "AMOUNT DUE:"
    totals_table.rows[3].cells[1].text = "R 3,450.00"
    
    # Make TOTAL row bold
    for cell in totals_table.rows[2].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(14)
    
    # Make AMOUNT DUE row bold and green
    for cell in totals_table.rows[3].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(16)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(39, 174, 96)  # Green
    
    for row in totals_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Payment terms
    p = doc.add_paragraph()
    run = p.add_run("PAYMENT TERMS:")
    run.font.bold = True
    doc.add_paragraph("Payment is due within 30 days of invoice date.")
    doc.add_paragraph("Please include invoice number on payment reference.")
    
    doc.add_paragraph()  # Spacing
    
    # Banking details
    p = doc.add_paragraph()
    run = p.add_run("BANKING DETAILS:")
    run.font.bold = True
    doc.add_paragraph("Bank: Your Bank Name")
    doc.add_paragraph("Account Name: Your Business Name")
    doc.add_paragraph("Account Number: XXXXXXXXXX")
    doc.add_paragraph("Branch Code: XXXXXX")
    
    doc.add_paragraph()  # Spacing
    
    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_p.add_run("Thank you for your business!")
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Save
    doc.save('/home/ubuntu/Operiva/artifacts/01-Full-Tax-Invoice-Template.docx')
    print("✓ Created 01-Full-Tax-Invoice-Template.docx")

def create_abridged_tax_invoice():
    """Create Abridged Tax Invoice template (for transactions ≤ R5,000)"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("TAX INVOICE (ABRIDGED)")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 42, 74)
    
    doc.add_paragraph()  # Spacing
    
    # Business details
    p = doc.add_paragraph()
    run = p.add_run("YOUR BUSINESS NAME")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph("Your Address, City, Province, Postal Code")
    doc.add_paragraph("Phone: +27 XX XXX XXXX | Email: your@email.com")
    doc.add_paragraph("VAT No: 4XXXXXXXXX")
    
    doc.add_paragraph()  # Spacing
    
    # Invoice details
    details_table = doc.add_table(rows=3, cols=2)
    details_table.rows[0].cells[0].text = "Invoice #:"
    details_table.rows[0].cells[1].text = "INV-001"
    details_table.rows[1].cells[0].text = "Date:"
    details_table.rows[1].cells[1].text = "2024/10/18"
    details_table.rows[2].cells[0].text = "Customer:"
    details_table.rows[2].cells[1].text = "Customer Name"
    
    doc.add_paragraph()  # Spacing
    
    # Items table (simpler for abridged)
    items_table = doc.add_table(rows=1, cols=3)
    items_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = items_table.rows[0].cells
    headers = ['DESCRIPTION', 'QTY', 'AMOUNT (incl. VAT)']
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1C2A4A')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Sample rows
    for i in range(2):
        row_cells = items_table.add_row().cells
        row_cells[0].text = f"Item {i+1} description"
        row_cells[1].text = "1"
        row_cells[2].text = "R 1,150.00"
    
    doc.add_paragraph()  # Spacing
    
    # Total
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("TOTAL (incl. VAT): R 2,300.00")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(39, 174, 96)
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = p.add_run("(VAT @ 15%: R 300.00)")
    run.font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Note about abridged invoice
    p = doc.add_paragraph()
    run = p.add_run("NOTE: ")
    run.font.bold = True
    run = p.add_run("This is an Abridged Tax Invoice valid for transactions up to R5,000 (incl. VAT).")
    run.font.italic = True
    run.font.size = Pt(9)
    
    # Save
    doc.save('/home/ubuntu/Operiva/artifacts/02-Abridged-Tax-Invoice-Template.docx')
    print("✓ Created 02-Abridged-Tax-Invoice-Template.docx")

def create_credit_note():
    """Create Credit Note template"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("CREDIT NOTE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 53, 69)  # Red for credit
    
    doc.add_paragraph()  # Spacing
    
    # Business details
    p = doc.add_paragraph()
    run = p.add_run("YOUR BUSINESS NAME")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph("Your Address, City, Province, Postal Code")
    doc.add_paragraph("Phone: +27 XX XXX XXXX | Email: your@email.com")
    doc.add_paragraph("VAT No: 4XXXXXXXXX")
    
    doc.add_paragraph()  # Spacing
    
    # Credit note details
    details_table = doc.add_table(rows=4, cols=2)
    details_table.rows[0].cells[0].text = "Credit Note #:"
    details_table.rows[0].cells[1].text = "CN-001"
    details_table.rows[1].cells[0].text = "Date:"
    details_table.rows[1].cells[1].text = "2024/10/18"
    details_table.rows[2].cells[0].text = "Original Invoice #:"
    details_table.rows[2].cells[1].text = "INV-001"
    details_table.rows[3].cells[0].text = "Customer:"
    details_table.rows[3].cells[1].text = "Customer Name"
    
    doc.add_paragraph()  # Spacing
    
    # Reason
    p = doc.add_paragraph()
    run = p.add_run("REASON FOR CREDIT:")
    run.font.bold = True
    doc.add_paragraph("Returned goods / Overcharge / Discount / Other: _________________")
    
    doc.add_paragraph()  # Spacing
    
    # Items table
    items_table = doc.add_table(rows=1, cols=4)
    items_table.style = 'Light Grid Accent 1'
    
    header_cells = items_table.rows[0].cells
    headers = ['DESCRIPTION', 'QTY', 'UNIT PRICE', 'AMOUNT']
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'DC3545')  # Red
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Sample row
    row_cells = items_table.add_row().cells
    row_cells[0].text = "Item description"
    row_cells[1].text = "1"
    row_cells[2].text = "R 1,000.00"
    row_cells[3].text = "R 1,000.00"
    
    doc.add_paragraph()  # Spacing
    
    # Totals
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    totals_table.rows[0].cells[0].text = "SUBTOTAL:"
    totals_table.rows[0].cells[1].text = "R 1,000.00"
    totals_table.rows[1].cells[0].text = "VAT (15%):"
    totals_table.rows[1].cells[1].text = "R 150.00"
    totals_table.rows[2].cells[0].text = "TOTAL CREDIT:"
    totals_table.rows[2].cells[1].text = "R 1,150.00"
    
    for cell in totals_table.rows[2].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(16)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
    
    for row in totals_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Note
    p = doc.add_paragraph()
    run = p.add_run("This credit note will be applied to your account. No refund will be issued unless specifically requested.")
    run.font.italic = True
    run.font.size = Pt(9)
    
    # Save
    doc.save('/home/ubuntu/Operiva/artifacts/03-Credit-Note-Template.docx')
    print("✓ Created 03-Credit-Note-Template.docx")

def create_quote_template():
    """Create Quotation template"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("QUOTATION")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 42, 74)
    
    doc.add_paragraph()  # Spacing
    
    # Business details
    p = doc.add_paragraph()
    run = p.add_run("YOUR BUSINESS NAME")
    run.font.size = Pt(14)
    run.font.bold = True
    
    doc.add_paragraph("Your Address, City, Province, Postal Code")
    doc.add_paragraph("Phone: +27 XX XXX XXXX | Email: your@email.com")
    doc.add_paragraph("VAT No: 4XXXXXXXXX")
    
    doc.add_paragraph()  # Spacing
    
    # Quote details
    details_table = doc.add_table(rows=4, cols=2)
    details_table.rows[0].cells[0].text = "Quote #:"
    details_table.rows[0].cells[1].text = "QUO-001"
    details_table.rows[1].cells[0].text = "Date:"
    details_table.rows[1].cells[1].text = "2024/10/18"
    details_table.rows[2].cells[0].text = "Valid Until:"
    details_table.rows[2].cells[1].text = "2024/11/18"
    details_table.rows[3].cells[0].text = "Customer:"
    details_table.rows[3].cells[1].text = "Customer Name"
    
    doc.add_paragraph()  # Spacing
    
    # Items table
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Light Grid Accent 1'
    
    header_cells = items_table.rows[0].cells
    headers = ['DESCRIPTION', 'QTY', 'UNIT PRICE', 'VAT', 'AMOUNT']
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '1C2A4A')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Sample rows
    for i in range(3):
        row_cells = items_table.add_row().cells
        row_cells[0].text = f"Item {i+1} description"
        row_cells[1].text = "1"
        row_cells[2].text = "R 1,000.00"
        row_cells[3].text = "R 150.00"
        row_cells[4].text = "R 1,150.00"
    
    doc.add_paragraph()  # Spacing
    
    # Totals
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    totals_table.rows[0].cells[0].text = "SUBTOTAL:"
    totals_table.rows[0].cells[1].text = "R 3,000.00"
    totals_table.rows[1].cells[0].text = "VAT (15%):"
    totals_table.rows[1].cells[1].text = "R 450.00"
    totals_table.rows[2].cells[0].text = "TOTAL:"
    totals_table.rows[2].cells[1].text = "R 3,450.00"
    
    for cell in totals_table.rows[2].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(16)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(39, 174, 96)
    
    for row in totals_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    doc.add_paragraph()  # Spacing
    
    # Terms
    p = doc.add_paragraph()
    run = p.add_run("TERMS & CONDITIONS:")
    run.font.bold = True
    doc.add_paragraph("• This quotation is valid for 30 days from the date above")
    doc.add_paragraph("• Prices include VAT at 15%")
    doc.add_paragraph("• Payment terms: 50% deposit, balance on completion")
    doc.add_paragraph("• Delivery: 14 working days from order confirmation")
    
    doc.add_paragraph()  # Spacing
    
    # Acceptance
    p = doc.add_paragraph()
    run = p.add_run("ACCEPTANCE:")
    run.font.bold = True
    doc.add_paragraph("I accept the above quotation:")
    doc.add_paragraph()
    doc.add_paragraph("Signature: _____________________  Date: _____________________")
    doc.add_paragraph("Name: _____________________")
    
    # Save
    doc.save('/home/ubuntu/Operiva/artifacts/04-Quotation-Template.docx')
    print("✓ Created 04-Quotation-Template.docx")

if __name__ == "__main__":
    print("Creating VAT-Aware Invoice Templates...")
    create_full_tax_invoice()
    create_abridged_tax_invoice()
    create_credit_note()
    create_quote_template()
    print("\n✅ All 4 invoice templates created successfully!")

